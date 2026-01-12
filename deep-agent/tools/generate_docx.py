"""DOCX generation tool."""
from langchain_core.tools import tool
from docx import Document




@tool
def generate_docx(content: str, filename: str = "output.docx") -> str:
    """Generates a DOCX file from text content."""
    doc = Document()
    doc.add_heading("Generated Document", 0)
    doc.add_paragraph(content)
    doc.save(filename)
    return f"DOCX saved: {filename}"
